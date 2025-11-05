# -*- coding: utf-8 -*-
"""
Sales Order PDF Generator – Dollar ($), o‘zbek lotin alifbosi, to‘liq ishlaydigan versiya
Fayl joyi: apps/cash_flow_app/cash_flow_app/custom_scripts/sales_order_pdf.py
"""
import frappe
from frappe import _
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime
import subprocess
from frappe.utils import flt, fmt_money


@frappe.whitelist()
def generate_contract_pdf(sales_order_name):
    """
    Sales Order uchun shartnoma PDF yaratish (USD, o‘zbek lotin)
    """
    try:
        # Sales Order ma'lumotlarini olish
        so = frappe.get_doc("Sales Order", sales_order_name)
        customer = frappe.get_doc("Customer", so.customer)
        company = frappe.get_doc("Company", so.company)

        # Template path
        site_path = frappe.get_site_path()
        template_path = os.path.join(site_path, "private", "files", "sales_order_template.docx")
        if not os.path.exists(template_path):
            return {
                "success": False,
                "message": "Template fayl topilmadi. Iltimos, sales_order_template.docx ni private/files papkasiga yuklab qo‘ying."
            }

        # Document yaratish
        doc = Document(template_path)

        # Ma'lumotlarni to'ldirish
        fill_contract_data(doc, so, customer, company)

        # DOCX saqlash
        output_dir = os.path.join(site_path, "private", "files", "contracts")
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        docx_filename = f"shartnoma_{so.name}_{timestamp}.docx"
        pdf_filename = f"shartnoma_{so.name}_{timestamp}.pdf"
        docx_path = os.path.join(output_dir, docx_filename)
        pdf_path = os.path.join(output_dir, pdf_filename)

        doc.save(docx_path)
        convert_to_pdf(docx_path, output_dir)

        # PDF ni Frappe ga qo'shish
        if os.path.exists(pdf_path):
            with open(pdf_path, 'rb') as f:
                file_doc = frappe.get_doc({
                    "doctype": "File",
                    "file_name": pdf_filename,
                    "attached_to_doctype": "Sales Order",
                    "attached_to_name": so.name,
                    "folder": "Home/Attachments",
                    "is_private": 1,
                    "content": f.read()
                })
                file_doc.save(ignore_permissions=True)
                frappe.db.commit()

            return {
                "success": True,
                "file_url": file_doc.file_url,
                "message": "Shartnoma PDF muvaffaqiyatli yaratildi!"
            }
        else:
            return {
                "success": False,
                "message": "PDF yaratishda xatolik yuz berdi"
            }

    except Exception as e:
        frappe.log_error(frappe.get_traceback(), "Sales Order PDF Xatosi")
        return {
            "success": False,
            "message": f"Xatolik: {str(e)}"
        }


def fill_contract_data(doc, so, customer, company):
    """
    Document ga ma'lumotlarni to'ldirish
    """
    # Sana formatlash
    date_str = so.transaction_date.strftime("%d.%m.%Y") if so.transaction_date else ""
    date_str_full = so.transaction_date.strftime("%d.%m.%Y") if so.transaction_date else ""

    # Mijoz ismi
    customer_full_name = customer.customer_name or customer.name

    # Asosiy matn almashtirish
    replacements = {
        "3000284959": so.name,
        "«13» Dekabr 2024": f"«{date_str}»",
        "13» Dekabr 2024": f"{date_str}»",
        "13.12.2024": date_str_full,
        "MCHJ\"TEXNOBARAKA\"": company.company_name,
        "ROVSHAN INOGAMXODJAYEV SAIDRAS": customer_full_name,
        "ROVSHAN INOGAMXODJAYEV SAIDRASULXUJAYEVICH": customer_full_name,
        "Алимбоев Сардор Дилхуш угли": company.get("custom_director_name") or "Rahbar",
    }

    # Paragraflarni yangilash
    for para in doc.paragraphs:
        for old, new in replacements.items():
            if old in para.text:
                for run in para.runs:
                    if old in run.text:
                        run.text = run.text.replace(old, new)

    # "сўм" so‘zini butunlay o‘chirish
    for para in doc.paragraphs:
        if "сўм" in para.text:
            for run in para.runs:
                run.text = run.text.replace("сўм", "").strip()

    # Jadvallarni to'ldirish
    tables = doc.tables
    if len(tables) >= 1:
        fill_items_table(tables[0], so)           # Mahsulotlar
    if len(tables) >= 2:
        fill_payment_table(tables[1], so)         # To‘lov shartlari
    if len(tables) >= 3:
        fill_parties_table(tables[2], customer, company)  # Tomonlar (1)
    if len(tables) >= 4:
        fill_acceptance_table(tables[3], so)      # Qabul dalolatnomasi
    if len(tables) >= 5:
        fill_schedule_table(tables[4], so)        # To‘lov jadvali
    if len(tables) >= 6:
        fill_parties_table(tables[5], customer, company)  # Imzolar (pastki)


def format_amount(amount):
    """
    Dollar format: 1 234.56 $   (2 kasr, nuqta, probel bilan)
    """
    if not amount:
        return "0.00 $"
    amount = flt(amount)
    # Frappe formatidan foydalanamiz (USD, 2 kasr)
    formatted = fmt_money(amount, precision=2, currency="USD")
    # "USD" → "$", probel bilan
    return formatted.replace("USD", "$").strip()


def clear_table_rows(table, keep_rows=1):
    """
    Jadvaldan headerdan tashqari barcha qatorlarni o‘chirish
    """
    while len(table.rows) > keep_rows:
        table._element.remove(table.rows[-1]._element)


def fill_items_table(table, so):
    """Mahsulotlar jadvali (Table 0)"""
    try:
        clear_table_rows(table, keep_rows=1)  # Header qoladi
        for idx, item in enumerate(so.items, 1):
            row = table.add_row()
            row.cells[0].text = str(idx)
            row.cells[1].text = item.item_name or item.item_code
            row.cells[2].text = "dona"
            row.cells[3].text = str(int(item.qty))
            row.cells[4].text = format_amount(item.rate)
            row.cells[5].text = format_amount(item.amount)
    except Exception as e:
        frappe.log_error(f"Items table error: {str(e)}", "PDF Generation")


def fill_payment_table(table, so):
    """To‘lov shartlari jadvali (Table 1)"""
    try:
        if len(table.rows) > 1:
            row = table.rows[1]
            row.cells[0].text = format_amount(so.total)
            row.cells[1].text = format_amount(so.grand_total)
            advance = flt(so.advance_paid or 0)
            row.cells[2].text = format_amount(advance)
            outstanding = flt(so.grand_total) - advance
            row.cells[3].text = format_amount(outstanding)
            # Muddat
            payment_period = frappe.db.get_value("Sales Order", so.name, "custom_payment_period")
            row.cells[4].text = payment_period or "1 oy"
    except Exception as e:
        frappe.log_error(f"Payment table error: {str(e)}", "PDF Generation")


def fill_parties_table(table, customer, company):
    """Tomonlar ma'lumotlari (Table 2 va 5)"""
    try:
        if len(table.rows) > 1:
            # Sotuvchi
            table.rows[1].cells[0].text = company.company_name
            # Xaridor
            customer_full_name = customer.customer_name or customer.name
            table.rows[1].cells[1].text = customer_full_name

        if len(table.rows) > 2:
            # Company manzili
            company_address = frappe.db.get_value(
                "Address", {"address_title": company.company_name}, "address_line1"
            ) or "Тошкент"

            # Customer manzili
            address_name = frappe.db.get_value(
                "Dynamic Link",
                {"link_doctype": "Customer", "link_name": customer.name, "is_primary_address": 1},
                "parent"
            )
            customer_address = ""
            if address_name:
                customer_address = frappe.db.get_value("Address", address_name, "address_line1") or ""

            table.rows[2].cells[0].text = f"Манзил: {company_address}"
            table.rows[2].cells[1].text = f"Доимий яшаш манзили: {customer_address}"

        # Bank va pasport ma'lumotlari (agar kerak bo‘lsa)
        if len(table.rows) > 3:
            # Misol: pastki qatorlarda bank ma'lumotlari
            pass

    except Exception as e:
        frappe.log_error(f"Parties table error: {str(e)}", "PDF Generation")


def fill_acceptance_table(table, so):
    """Qabul-topshirish dalolatnomasi (Table 3)"""
    try:
        clear_table_rows(table, keep_rows=1)  # Header + jami qatori qoladi
        for idx, item in enumerate(so.items, 1):
            row = table.add_row()
            row.cells[0].text = str(idx)
            row.cells[1].text = item.item_name or item.item_code
            row.cells[2].text = "dona"
            row.cells[3].text = str(int(item.qty))
            row.cells[4].text = format_amount(item.rate)
            row.cells[5].text = "0 $"

        # Jami qatori (agar jadvalda bo‘lsa)
        if len(table.rows) > 1:
            last_row = table.rows[-1]
            if len(last_row.cells) >= 5:
                last_row.cells[4].text = format_amount(so.grand_total)

    except Exception as e:
        frappe.log_error(f"Acceptance table error: {str(e)}", "PDF Generation")


def fill_schedule_table(table, so):
    """To‘lov jadvali (Table 4)"""
    try:
        clear_table_rows(table, keep_rows=1)  # Header qoladi

        payment_schedule = frappe.get_all(
            "Payment Schedule",
            filters={"parent": so.name, "parenttype": "Sales Order"},
            fields=["due_date", "payment_amount"],
            order_by="due_date"
        )

        if payment_schedule:
            for idx, sch in enumerate(payment_schedule, 1):
                row = table.add_row()
                row.cells[0].text = str(idx)
                row.cells[1].text = sch.due_date.strftime("%d.%m.%Y") if sch.due_date else ""
                row.cells[2].text = format_amount(sch.payment_amount)
        else:
            # Default: bitta to‘lov
            row = table.add_row()
            row.cells[0].text = "1"
            delivery_date = so.delivery_date or so.transaction_date
            row.cells[1].text = delivery_date.strftime("%d.%m.%Y") if delivery_date else ""
            row.cells[2].text = format_amount(so.grand_total)

    except Exception as e:
        frappe.log_error(f"Schedule table error: {str(e)}", "PDF Generation")


def convert_to_pdf(docx_path, output_dir):
    """LibreOffice orqali DOCX → PDF"""
    try:
        subprocess.run([
            "libreoffice",
            "--headless",
            "--convert-to", "pdf",
            "--outdir", output_dir,
            docx_path
        ], check=True, timeout=30)
    except subprocess.TimeoutExpired:
        frappe.throw("PDF yaratish vaqti tugadi. Iltimos, qaytadan urinib ko‘ring.")
    except FileNotFoundError:
        frappe.throw("LibreOffice o‘rnatilmagan. Server administratoriga murojaat qiling.")
    except Exception as e:
        frappe.throw(f"PDF yaratishda xatolik: {str(e)}")